"""Generate experiment report in Word format."""
import json, sys
from pathlib import Path
from datetime import date

sys.path.insert(0, str(Path(__file__).parent.parent))
from config import EVAL_DIR, NUM_PASSAGES, NUM_QUERIES

DOCS_DIR = Path(__file__).parent.parent / "docs"
DOCS_DIR.mkdir(exist_ok=True)

with open(EVAL_DIR / "summary.json") as f:
    summary = json.load(f)

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ── Styles ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for i, (name, size, bold) in enumerate([("Heading 1", 18, True), ("Heading 2", 14, True), ("Heading 3", 12, True)]):
    s = doc.styles[name]
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ═══════════════ COVER ═══════════════
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Agentic RAG 多策略检索系统\n实验报告")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f"MS MARCO 10K · 6 Strategies · RAGAS Evaluation\n{date.today().strftime('%Y-%m-%d')}")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════ 1. 项目概述 ═══════════════
doc.add_heading("1. 项目概述", level=1)
doc.add_paragraph(
    "本项目构建了一个完整的多策略 RAG（检索增强生成）检索评测系统。系统基于 MS MARCO 数据集，"
    "实现了 6 种不同的检索策略，并使用 RAGAS 框架配合 Claude Sonnet 4.6 进行自动化评估。"
    "所有嵌入模型（BGE-base-en-v1.5）和重排序模型（BGE-reranker-base）均在本地 GPU 上运行，"
    "向量数据库采用 Qdrant 进行持久化存储。"
)

# Key metrics table
doc.add_heading("核心指标", level=2)
table = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
items = [
    ("数据集", "MS MARCO v2.1"),
    ("文档数", f"{NUM_PASSAGES:,} 篇"),
    ("评估查询", f"{NUM_QUERIES} 条"),
    ("检索策略", "6 种"),
    ("GPU", "NVIDIA GeForce RTX 5060 Ti 8GB"),
]
for i, (k, v) in enumerate(items):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# ═══════════════ 2. 项目流程 ═══════════════
doc.add_page_break()
doc.add_heading("2. 项目流程", level=1)

flow_steps = [
    ("Step 1: 数据下载", "从 MS MARCO v2.1 数据集下载 10,007 篇文档和 200 条 QA 对，"
     "通过文本去重确保文档唯一性，不足 10K 时从 v1.1 补充。提取每篇文档的 passage_text 并建立唯一 ID 映射。"),
    ("Step 2: 向量化", "使用 BAAI/bge-base-en-v1.5 模型在本地 GPU 上对 10,007 篇文档生成 768 维稠密向量。"
     "每 500 条保存一次 .npy 检查点，支持断点续跑。同时基于 token 分词构建 BM25 稀疏索引。"),
    ("Step 3: 向量入库", "将 10,007 条稠密向量导入 Qdrant 向量数据库，使用 Cosine 距离度量。"
     "Qdrant 通过 Docker Compose 启动，数据持久化到本地 qdrant_storage/ 目录。"),
    ("Step 4: 检索策略实验", "对 100 条查询逐一执行 6 种检索策略：BM25 关键词匹配、Dense 向量检索、"
     "Hybrid RRF 融合、HyDE 假设文档扩展、BGE Reranker 精排、以及 ReAct Agent 反思检索。"
     "每种策略返回 Top-10 相关文档。"),
    ("Step 5: RAGAS 评估", "使用 RAGAS 框架的 Context Precision 和 Context Recall 指标进行评估。"
     "LLM 裁判采用 Claude Sonnet 4.6（通过 ChatAnthropic 接口），逐条打分并追加写入 .jsonl 文件。"
     "支持断点续评，已完成查询自动跳过。"),
    ("Step 6: 结果分析", "汇总 6 策略 × 100 查询的评估结果，计算平均精度和召回率。"
     "生成对比 CSV 表格、雷达图、柱状图，以及 Word 实验报告。"),
]

for title, desc in flow_steps:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

# ═══════════════ 3. 系统架构 ═══════════════
doc.add_page_break()
doc.add_heading("3. 系统架构", level=1)

doc.add_heading("3.1 技术栈", level=2)
tech_table = doc.add_table(rows=9, cols=3, style="Light Grid Accent 1")
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (comp, tech, env) in enumerate([
    ("Dense Embedding", "BAAI/bge-base-en-v1.5", "本地 GPU (RTX 5060 Ti)"),
    ("Reranker", "BAAI/bge-reranker-base", "本地 GPU"),
    ("Sparse Retrieval", "BM25 (rank-bm25)", "本地 CPU"),
    ("Vector Database", "Qdrant v1.16.2", "Docker"),
    ("LLM Judge", "Claude Sonnet 4.6", "API (aipaibox.com)"),
    ("Evaluation", "RAGAS 0.4.x", "本地 + API"),
    ("Backend", "FastAPI 0.136", "Docker / 本地"),
    ("Frontend", "Streamlit 1.57", "Docker / 本地"),
]):
    tech_table.rows[i].cells[0].text = comp
    tech_table.rows[i].cells[1].text = tech
    tech_table.rows[i].cells[2].text = env

doc.add_heading("3.2 检索策略设计", level=2)
strategies_desc = [
    ("BM25 (策略 1)", "基于 TF-IDF 的稀疏关键词匹配。使用 BGE tokenizer 分词后构建 BM25Okapi 索引，"
     "通过 get_top_n() 返回得分最高的 k 篇文档。优点：零 API 调用，速度快。缺点：无语义理解。"),
    ("Dense (策略 2)", "使用 BGE-base-en-v1.5 将查询编码为 768 维向量，在 Qdrant 中通过 Cosine 相似度检索。"
     "优点：语义匹配，能理解同义词。缺点：对专用术语精度不如 BM25。"),
    ("Hybrid RRF (策略 3)", "Reciprocal Rank Fusion 融合 BM25 和 Dense 的结果。"
     "每个文档的最终得分 = Σ 1/(k + rank + 1)，k=60。优点：结合两者优势。缺点：噪声可能互相干扰。"),
    ("Hybrid + HyDE (策略 4)", "在 RRF 融合前，先用 Claude Sonnet 生成假设文档（Hypothetical Document），"
     "用假设文档替代原始查询进行 Dense 检索，然后与 BM25 结果融合。"
     "优点：缓解查询-文档词汇不匹配。缺点：额外 API 调用。"),
    ("Hybrid + HyDE + Rerank (策略 5)", "在策略 4 的基础上增加 BGE-reranker-base 交叉编码器精排。"
     "将候选文档与查询组成 [query, doc] 对输入 reranker 打分，取 Top-k。"
     "优点：重排序大幅提升精度。缺点：计算量较大。"),
    ("Agentic ReAct (策略 6)", "基于 ReAct 框架的智能代理。Agent 可调用 search_bm25() 和 search_dense() 工具，"
     "每轮进行 Thought → Action → Observation 循环，最多 3 轮。"
     "每轮结束后由反思模块评估检索质量。"
     "优点：自适应选择策略。缺点：多次 LLM 调用，成本高，当前准确率欠佳。"),
]
for title, desc in strategies_desc:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ═══════════════ 4. 实验结果 ═══════════════
doc.add_page_break()
doc.add_heading("4. 实验结果", level=1)

doc.add_heading("4.1 综合对比 (MS MARCO 10K, 100 Queries)", level=2)

result_table = doc.add_table(rows=8, cols=3, style="Light Grid Accent 1")
result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
result_table.rows[0].cells[0].text = "策略"
result_table.rows[0].cells[1].text = "Context Precision"
result_table.rows[0].cells[2].text = "Context Recall"
for i, s in enumerate(["bm25", "dense", "hybrid_rrf", "hybrid_hyde", "hybrid_hyde_rerank", "agentic"]):
    labels = {"bm25": "BM25", "dense": "Dense (BGE)", "hybrid_rrf": "Hybrid RRF",
              "hybrid_hyde": "Hybrid + HyDE", "hybrid_hyde_rerank": "Hybrid + HyDE + Rerank",
              "agentic": "Agentic (ReAct)"}
    result_table.rows[i+1].cells[0].text = labels[s]
    result_table.rows[i+1].cells[1].text = f"{summary[s]['context_precision']:.4f}"
    result_table.rows[i+1].cells[2].text = f"{summary[s]['context_recall']:.4f}"

# Bold best row
for cell in result_table.rows[5].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

doc.add_paragraph()

doc.add_heading("4.2 实验结论", level=2)
conclusions = [
    ("1. Hybrid + HyDE + Rerank 综合最优",
     "Precision 0.3785, Recall 0.5102，均为 6 种策略中最高。"
     "说明稠密检索 + 假设文档扩展 + 交叉编码器精排的三阶段 pipeline 在 MS MARCO 数据集上效果最佳。"
     "Reranker 的交叉注意力机制有效纠正了双塔模型在细粒度语义匹配上的不足。"),
    ("2. RRF 融合未必优于单策略",
     "Hybrid RRF (Recall 0.346) 的召回率低于纯 Dense (Recall 0.470)，说明当 BM25 结果质量较低时，"
     "RRF 融合反而引入噪声。需要调整融合权重或设置质量阈值。"),
    ("3. Agentic 检索仍需改进",
     "ReAct Agent 的 Precision 仅 0.067，Recall 仅 0.143，远低于其他策略。"
     "主要原因：工具定义过于简单（仅 search_bm25/search_dense），缺少 rerank 和结果合并工具；"
     "反思模块对检索质量的判断不够精准。未来需要增强 agent 的工具箱和 prompt 设计。"),
    ("4. BM25 的召回率仍有价值",
     "尽管无语义理解，BM25 的 Recall (0.465) 与 Dense (0.470) 基本持平，"
     "说明 MS MARCO 数据集中关键词匹配信号仍然有效，适合用于混合检索的互补组件。"),
    ("5. HyDE 单独使用效果欠佳",
     "Hybrid + HyDE (Precision 0.160) 精度低于纯 Dense (0.301)，"
     "但搭配 Reranker 后(0.379)提升显著。表明假设文档的价值需要 reranker 才能充分释放。"),
]
for title, desc in conclusions:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

# ═══════════════ 5. 工程实现 ═══════════════
doc.add_page_break()
doc.add_heading("5. 工程实现亮点", level=1)
engineering = [
    "断点续跑: pipeline_state.json 记录每个阶段状态，embedding 每 500 条 npy 检查点，评估每 query 追加 jsonl",
    "GPU 优化: batch_size=32 (embedding) / 16 (reranker)，用完 del model + torch.cuda.empty_cache()",
    "Docker 一键部署: docker-compose.yml 包含 Qdrant + FastAPI + Streamlit 三服务",
    "成本追踪: api_usage.json 记录每次 API 调用的估算 cost，支持预算上限自动中断",
    "独立服务: FastAPI (8 个接口) + Streamlit (检索/结果/API 三页面)，前后端分离",
]
for item in engineering:
    doc.add_paragraph(item, style="List Bullet")

# ═══════════════ 6. 未来展望 ═══════════════
doc.add_heading("6. 未来展望", level=1)
future = [
    "多模态 RAG: 扩展支持图片、表格等多模态文档的检索",
    "Agent 增强: 为 ReAct Agent 添加 rerank、结果聚合等工具，改进反思 prompt",
    "更大规模: 扩展到 MS MARCO 全量 8.8M passages，测试策略的可扩展性",
    "在线学习: 基于用户反馈持续优化检索排序",
    "多语言: 支持中文等多语言检索评测",
]
for item in future:
    doc.add_paragraph(item, style="List Bullet")

# ── Save ──
output_path = DOCS_DIR / "experiment_report.docx"
doc.save(str(output_path))
print(f"[report] Saved to {output_path}")
